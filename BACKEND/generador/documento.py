"""
Módulo para generación de documentos Word
Genera el informe completo en formato DOCX
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .graficas import crear_grafica_circular
from .analizador import analizar_columna, generar_analisis_resultados, generar_oportunidades_mejora
from .utils import convertir_a_png, ENCABEZADO, PIE_PAGINA, COLUMNAS_EXCLUIR


def _agregar_encabezado_pie(doc, directorio_trabajo='.'):
    """
    Agrega encabezado y pie de página al documento
    
    Args:
        doc: Documento de python-docx
        directorio_trabajo: Directorio donde buscar las imágenes
    """
    # CONFIGURAR MÁRGENES DE LA PÁGINA
    section = doc.sections[0]
    
    # Márgenes del contenido según especificaciones exactas
    section.top_margin = Inches(1)          # 2.54 cm
    section.bottom_margin = Inches(1)       # 2.54 cm
    section.left_margin = Inches(1.248)     # 3.17 cm
    section.right_margin = Inches(1.248)    # 3.17 cm
    
    # Encabezado y pie: 0 cm desde los bordes
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)
    
    header = section.header
    
    # Buscar imagen de encabezado en múltiples ubicaciones
    imagen_encabezado = None
    
    # 1. Primero buscar en directorio de trabajo (para imágenes subidas por el usuario)
    for ext in ['.png', '.jpg', '.jpeg', '.PNG', '.JPG', '.JPEG']:
        archivo = os.path.join(directorio_trabajo, f'encabezado{ext}')
        if os.path.exists(archivo):
            imagen_encabezado = archivo
            break
    
    # 2. Si no se encontró, buscar en carpeta imagenes/ (imágenes estáticas)
    if not imagen_encabezado:
        # Obtener ruta base del proyecto (donde está app.py)
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        imagenes_dir = os.path.join(base_dir, 'imagenes')
        
        for ext in ['.png', '.jpg', '.jpeg', '.PNG', '.JPG', '.JPEG']:
            archivo = os.path.join(imagenes_dir, f'encabezado{ext}')
            if os.path.exists(archivo):
                imagen_encabezado = archivo
                print(f"  ℹ️  Usando imagen estática: {archivo}")
                break
    
    # Si existe imagen de encabezado, agregarla
    if imagen_encabezado:
        imagen_encabezado = convertir_a_png(imagen_encabezado)
        
        p_img = header.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_img.paragraph_format.left_indent = Inches(0)
        p_img.paragraph_format.right_indent = Inches(0)
        p_img.paragraph_format.space_before = Pt(0)
        p_img.paragraph_format.space_after = Pt(0)
        
        run_img = p_img.add_run()
        run_img.add_picture(imagen_encabezado, width=Inches(8.5))
    else:
        # Usar texto por defecto
        p1 = header.paragraphs[0]
        p1.text = ENCABEZADO['linea1']
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.runs[0]
        run1.font.size = Pt(11)
        run1.font.bold = True
        
        p2 = header.add_paragraph()
        p2.text = ENCABEZADO['linea2']
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.runs[0]
        run2.font.size = Pt(11)
        run2.font.bold = True
        
        p3 = header.add_paragraph()
        p3.text = ENCABEZADO['linea3']
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.runs[0]
        run3.font.size = Pt(9)
        run3.font.bold = True
    
    # PIE DE PÁGINA
    footer = section.footer
    
    # Buscar imagen de pie
    imagen_pie = None
    
    # 1. Primero buscar en directorio de trabajo
    for ext in ['.png', '.jpg', '.jpeg', '.PNG', '.JPG', '.JPEG']:
        archivo = os.path.join(directorio_trabajo, f'pie{ext}')
        if os.path.exists(archivo):
            imagen_pie = archivo
            break
    
    # 2. Si no se encontró, buscar en carpeta imagenes/
    if not imagen_pie:
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        imagenes_dir = os.path.join(base_dir, 'imagenes')
        
        for ext in ['.png', '.jpg', '.jpeg', '.PNG', '.JPG', '.JPEG']:
            archivo = os.path.join(imagenes_dir, f'pie{ext}')
            if os.path.exists(archivo):
                imagen_pie = archivo
                print(f"  ℹ️  Usando imagen estática: {archivo}")
                break
    
    if imagen_pie:
        imagen_pie = convertir_a_png(imagen_pie)
        
        pf_img = footer.paragraphs[0]
        pf_img.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf_img.paragraph_format.left_indent = Inches(0)
        pf_img.paragraph_format.right_indent = Inches(0)
        pf_img.paragraph_format.space_before = Pt(0)
        pf_img.paragraph_format.space_after = Pt(0)
        
        run_img_f = pf_img.add_run()
        run_img_f.add_picture(imagen_pie, width=Inches(8.5))
    else:
        # Usar texto por defecto
        pf1 = footer.paragraphs[0]
        pf1.text = PIE_PAGINA['linea1']
        pf1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        runf1 = pf1.runs[0]
        runf1.font.size = Pt(9)
        runf1.font.bold = True
        
        pf2 = footer.add_paragraph()
        pf2.text = PIE_PAGINA['linea2']
        pf2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        runf2 = pf2.runs[0]
        runf2.font.size = Pt(9)
        runf2.font.bold = True
        
        pf3 = footer.add_paragraph()
        pf3.text = PIE_PAGINA['linea3']
        pf3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        runf3 = pf3.runs[0]
        runf3.font.size = Pt(9)
        runf3.font.bold = True


def _convertir_valor_texto(valor):
    """
    Convierte valores numéricos a texto descriptivo
    
    Args:
        valor: Valor a convertir
        
    Returns:
        str: Texto descriptivo
    """
    valor_str = str(valor).strip()
    
    if valor_str == '5':
        return 'muy satisfactorio'
    elif valor_str == '4':
        return 'satisfactorio'
    elif valor_str == '3':
        return 'aceptable'
    elif valor_str == '2':
        return 'insatisfactorio'
    elif valor_str == '1':
        return 'muy insatisfactorio'
    else:
        return valor_str


def _generar_texto_resultado(items):
    """
    Genera el texto descriptivo de los resultados
    
    Args:
        items: Lista de tuplas (valor, porcentaje)
        
    Returns:
        str: Texto descriptivo
    """
    # Detectar si son valores numéricos
    es_numerica = all(item[0].strip() in ['1', '2', '3', '4', '5'] for item in items)
    
    # Detectar tipo de respuesta
    primera_opcion = items[0][0].lower()
    es_satisfaccion = any(palabra in primera_opcion for palabra in ['satisfecho', 'satisfactorio', 'satisfactoria'])
    es_calificacion = any(palabra in primera_opcion for palabra in ['bueno', 'buena', 'malo', 'mala', 'regular', 'excelente'])
    
    if len(items) == 1:
        valor_texto = _convertir_valor_texto(items[0][0]) if es_numerica else items[0][0]
        
        if es_numerica:
            return f"El {items[0][1]}% da una calificación de {valor_texto}."
        elif es_satisfaccion:
            return f"El {items[0][1]}% dio una respuesta de {valor_texto}."
        elif es_calificacion:
            return f"El {items[0][1]}% dan una calificación de {valor_texto} a la pregunta."
        else:
            return f"El {items[0][1]}% respondieron {valor_texto}."
            
    elif len(items) == 2:
        if es_numerica:
            valor1 = _convertir_valor_texto(items[0][0])
            valor2 = _convertir_valor_texto(items[1][0])
            return f"Las respuesta es un {items[0][1]}% da una calificación de {valor1} y un {items[1][1]}% da una calificación de {valor2}."
        elif es_satisfaccion:
            return f"El {items[0][1]}% están {items[0][0]} y el {items[1][1]}% {items[1][0]}."
        elif es_calificacion:
            return f"El {items[0][1]}% dan una calificación de {items[0][0]} a la pregunta y el {items[1][1]}% da una calificación de {items[1][0]}."
        else:
            return f"El {items[0][1]}% dio una respuesta de {items[0][0]} y el {items[1][1]}% de {items[1][0]}."
            
    elif len(items) == 3:
        if es_numerica:
            valor1 = _convertir_valor_texto(items[0][0])
            valor2 = _convertir_valor_texto(items[1][0])
            valor3 = _convertir_valor_texto(items[2][0])
            return f"El {items[0][1]}% da una calificación de {valor1}, el {items[1][1]}% de {valor2} y el {items[2][1]}% de {valor3}."
        elif es_satisfaccion:
            return f"El {items[0][1]}% están {items[0][0]}, el {items[1][1]}% {items[1][0]} y el {items[2][1]}% {items[2][0]}."
        elif es_calificacion:
            return f"El {items[0][1]}% dan una calificación de {items[0][0]}, el {items[1][1]}% de {items[1][0]} y el {items[2][1]}% de {items[2][0]}."
        else:
            return f"El {items[0][1]}% respondieron {items[0][0]}, el {items[1][1]}% {items[1][0]} y el {items[2][1]}% {items[2][0]}."
            
    else:
        # Más de tres opciones
        partes = []
        for i, (key, value) in enumerate(items):
            valor_texto = _convertir_valor_texto(key) if es_numerica else key
            
            if i == 0:
                partes.append(f"El {value}% da una calificación de {valor_texto}")
            elif i == len(items) - 1:
                partes.append(f"y el {value}% de {valor_texto}")
            else:
                partes.append(f"el {value}% de {valor_texto}")
        return f"{', '.join(partes)}."


def generar_informe_word(archivo_excel, archivo_salida, nombre_uds='', directorio_trabajo='.'):
    """
    Genera el informe completo en formato Word
    
    Args:
        archivo_excel: Ruta del archivo Excel con los datos
        archivo_salida: Ruta del archivo de salida .docx
        nombre_uds: Nombre de la Unidad de Servicio (opcional)
        directorio_trabajo: Directorio de trabajo para las imágenes
        
    Returns:
        dict: Diccionario con información del proceso
    """
    print(f"📊 Leyendo archivo: {archivo_excel}")
    
    # Extraer nombre del UDS si no se proporciona
    if not nombre_uds:
        nombre_archivo = os.path.basename(archivo_excel)
        nombre_sin_ext = os.path.splitext(nombre_archivo)[0]
        nombre_uds = nombre_sin_ext.replace('_', ' ').title()
        print(f"📝 Nombre UDS detectado: {nombre_uds}")
    
    # Leer Excel
    df = pd.read_excel(archivo_excel)
    total_respuestas = len(df)
    print(f"✓ {total_respuestas} respuestas encontradas")
    print(f"✓ {len(df.columns)} columnas detectadas")
    
    # Crear documento
    doc = Document()
    
    # Agregar encabezado y pie
    _agregar_encabezado_pie(doc, directorio_trabajo)
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # INTRODUCCIÓN
    p_intro_titulo = doc.add_paragraph()
    run = p_intro_titulo.add_run('Introducción')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    p_intro_titulo.paragraph_format.space_after = Pt(6)
    
    p_intro_texto = doc.add_paragraph()
    texto_intro = (f"Dentro del marco de las obligaciones contractuales SIGE establecidas entre "
                   f"EL INSTITUTO COLOMBIANO DE BIENESTAR FAMILIAR ICBF y la ASOCIACION DE PADRES "
                   f"DE FAMILIA DEL HOGAR INFANTIL GUATAPURI (UDS) {nombre_uds.upper()} se establecer el de realizar "
                   f"una encuesta que permita saber el nivel de satisfacción de los usuarios respecto "
                   f"al servicio prestado el siguiente documento muestra la metodología, los resultados, "
                   f"el análisis de los mismos y unas posibles oportunidades de mejora.")
    run = p_intro_texto.add_run(texto_intro)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    p_intro_texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_intro_texto.paragraph_format.space_after = Pt(11)
    
    # METODOLOGÍA
    p_metod_titulo = doc.add_paragraph()
    run = p_metod_titulo.add_run('Metodología')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    p_metod_titulo.paragraph_format.space_after = Pt(6)
    
    p_metod_texto = doc.add_paragraph()
    texto_metod = (f"El primer paso de la metodología consistió en la elaboración de una encuesta "
                   f"(lista de preguntas con calificación) que permitiría saber el nivel de satisfacción "
                   f"de los usuarios de cada uds (para este caso fue la uds {nombre_uds.title()}) respectos a los "
                   f"distintos ítems de calificación del servicio estas preguntas se establecieron en un "
                   f"orden de 1 a 5 donde uno es muy malo y 5 muy bueno, y algunas de si o no una vez "
                   f"establecidas estas preguntas se estableció un formulario tipo GOOGLE y se vinculó al "
                   f"correo gerenciamasisosas@gmail.com, antes del inicio del encuentro se le explico a los "
                   f"50 usuarios de la uds la importancia del diligenciamiento de la encuesta, por medios "
                   f"electrónicos se le envió a los usuarios la encuesta a diligenciar por lineamientos del "
                   f"ICBF se establece un mínimo del 20% de la población como muestra, para este caso se "
                   f"lograron diligenciar {total_respuestas} encuestas, una vez diligenciadas se procederá "
                   f"a realizar las fase de RESULTADOS, ANALISIS DE RESULTADOS Y POSIBLES OPORTUNIDADES DE MEJORA.")
    run = p_metod_texto.add_run(texto_metod)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    p_metod_texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_metod_texto.paragraph_format.space_after = Pt(11)
    
    # RESULTADOS (Página 2)
    doc.add_page_break()
    
    p_result_titulo = doc.add_paragraph()
    run = p_result_titulo.add_run('Resultados')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    p_result_titulo.paragraph_format.space_after = Pt(12)
    
    # Procesar cada pregunta
    print("\n📈 Generando análisis y gráficas...")
    
    resultados_todas_preguntas = []
    contador_preguntas = 0
    
    for columna in df.columns:
        # Verificar si la columna debe excluirse
        columna_lower = columna.lower().strip()
        if any(excluir in columna_lower for excluir in COLUMNAS_EXCLUIR):
            print(f"  ⏭️  Omitiendo: {columna}")
            continue
        
        print(f"  Procesando: {columna}")
        
        resultado = analizar_columna(df, columna)
        
        if resultado is None:
            print(f"  ⚠️  Columna vacía, omitiendo")
            continue
        
        resultados_todas_preguntas.append(resultado)
        
        # Texto pregunta
        p_pregunta = doc.add_paragraph()
        texto_pregunta = f'Ante la pregunta "{resultado["pregunta"]}" Los resultados se muestran en la siguiente gráfica.'
        run = p_pregunta.add_run(texto_pregunta)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        p_pregunta.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_pregunta.paragraph_format.space_after = Pt(6)
        
        # Gráfica
        img_buffer = crear_grafica_circular(
            resultado['porcentajes_exactos'], 
            resultado['pregunta']
        )
        
        p_grafica = doc.add_paragraph()
        p_grafica.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_grafica.add_run()
        run.add_picture(img_buffer, width=Inches(3.65))
        p_grafica.paragraph_format.space_after = Pt(6)
        
        # Texto resultados
        items = list(resultado['porcentajes'].items())
        texto_resultado = _generar_texto_resultado(items)
        
        p_resultados = doc.add_paragraph()
        run = p_resultados.add_run(texto_resultado)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        p_resultados.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_resultados.paragraph_format.space_after = Pt(12)
        
        # Salto de página cada 2 preguntas
        contador_preguntas += 1
        if contador_preguntas % 2 == 0:
            doc.add_page_break()
    
    # ANÁLISIS DE RESULTADOS
    print("\n📊 Generando análisis de resultados...")
    
    if contador_preguntas % 2 != 0:
        doc.add_page_break()
    
    p_analisis_titulo = doc.add_paragraph()
    run = p_analisis_titulo.add_run('Análisis de resultados')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    p_analisis_titulo.paragraph_format.space_before = Pt(12)
    p_analisis_titulo.paragraph_format.space_after = Pt(6)
    
    texto_analisis = generar_analisis_resultados(resultados_todas_preguntas, nombre_uds)
    
    p_analisis_texto = doc.add_paragraph()
    run = p_analisis_texto.add_run(texto_analisis)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    p_analisis_texto.paragraph_format.space_after = Pt(12)
    p_analisis_texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # OPORTUNIDADES DE MEJORA
    print("💡 Generando oportunidades de mejora...")
    
    p_oportunidades_titulo = doc.add_paragraph()
    run = p_oportunidades_titulo.add_run('Posibles oportunidades de mejora')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    p_oportunidades_titulo.paragraph_format.space_before = Pt(12)
    p_oportunidades_titulo.paragraph_format.space_after = Pt(6)
    
    oportunidades = generar_oportunidades_mejora(resultados_todas_preguntas)
    
    for oportunidad in oportunidades:
        p_oportunidad = doc.add_paragraph()
        p_oportunidad.paragraph_format.left_indent = Inches(0.25)
        run = p_oportunidad.add_run(f"·       {oportunidad}")
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        p_oportunidad.paragraph_format.space_after = Pt(6)
        p_oportunidad.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Guardar documento
    doc.save(archivo_salida)
    
    print(f"\n✅ Informe generado exitosamente: {archivo_salida}")
    
    return {
        'success': True,
        'archivo_salida': archivo_salida,
        'nombre_uds': nombre_uds,
        'total_respuestas': total_respuestas,
        'total_preguntas': len(resultados_todas_preguntas)
    }